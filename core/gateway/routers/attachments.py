"""
attachments.py — Adjuntar documentos al chat.
Extrae el TEXTO de un documento (PDF, Word, texto plano, código) para inyectarlo
como contexto del mensaje. No guarda el archivo: los modelos solo necesitan el texto.

Las imágenes NO pasan por aquí: el chat las manda a /api/vision/describe, donde un
modelo multimodal las describe, y esa descripción se adjunta como texto. El audio
no tiene camino todavía — el gateway no transcribe.
"""
from __future__ import annotations

import io
import logging
from typing import Any

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from core.security.auth import cookie_auth_required

logger = logging.getLogger("lixbon.attachments")
router = APIRouter()

MAX_FILE_BYTES = 5 * 1024 * 1024        # 5 MB
MAX_TEXT_CHARS = 12_000                 # ~3k tokens: cabe en el contexto de modelos pequeños

# Extensiones de texto/código que leemos directamente
_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".yaml", ".yml", ".xml",
    ".html", ".css", ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".c", ".cpp",
    ".h", ".hpp", ".cs", ".go", ".rs", ".rb", ".php", ".sh", ".sql", ".ini",
    ".toml", ".log", ".rtf", ".tex",
}


def _ext(name: str) -> str:
    dot = name.rfind(".")
    return name[dot:].lower() if dot >= 0 else ""


def _extract_docx(data: bytes) -> str:
    """Texto de un .docx: párrafos y, además, el contenido de las tablas, que en
    un documento de trabajo suele ser justo lo que importa."""
    from docx import Document  # dependencia opcional: si falta, el .docx cae en «no soportado»

    doc = Document(io.BytesIO(data))
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))
    return '\n'.join(partes)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


@router.post("/api/attachments")
async def upload_attachment(
    file: UploadFile = File(...),
    _user: dict[str, Any] = Depends(cookie_auth_required),
):
    """Sube un documento y devuelve su texto extraído (para adjuntar al mensaje)."""
    data = await file.read()
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail={
            "code": "too_large", "message": "El archivo supera el límite de 5 MB.",
        })

    name = file.filename or "documento"
    ext = _ext(name)

    if ext == ".pdf" or (file.content_type or "").endswith("pdf"):
        try:
            text = _extract_pdf(data)
        except Exception as exc:
            logger.warning(f"PDF ilegible: {exc}")
            raise HTTPException(status_code=422, detail={
                "code": "pdf_unreadable", "message": "No se pudo leer el texto de ese PDF.",
            })
    elif ext == ".docx":
        try:
            text = _extract_docx(data)
        except ImportError:
            raise HTTPException(status_code=415, detail={
                "code": "unsupported",
                "message": "Este servidor no puede leer archivos .docx todavía.",
            })
        except Exception as exc:
            logger.warning(f"DOCX ilegible: {exc}")
            raise HTTPException(status_code=422, detail={
                "code": "docx_unreadable", "message": "No se pudo leer el texto de ese documento.",
            })
    elif ext in _TEXT_EXTS or (file.content_type or "").startswith("text/"):
        try:
            text = data.decode("utf-8", errors="replace")
        except Exception:
            raise HTTPException(status_code=422, detail={
                "code": "decode_error", "message": "No se pudo leer el texto del archivo.",
            })
    else:
        raise HTTPException(status_code=415, detail={
            "code": "unsupported",
            "message": "Solo se admiten documentos (PDF, Word, texto y código) e "
                       "imágenes. El audio todavía no: usa el micrófono para dictar.",
        })

    text = text.strip()
    if not text:
        raise HTTPException(status_code=422, detail={
            "code": "empty", "message": "El documento no tiene texto legible.",
        })

    truncated = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS]

    return {
        "filename": name,
        "text": text,
        "chars": len(text),
        "truncated": truncated,
    }
